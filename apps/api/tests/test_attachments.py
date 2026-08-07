import io
from unittest.mock import MagicMock, patch

import httpx
import pytest
from docx import Document
from PIL import Image

from app.content_pipeline.attachments import (
    AttachmentTooLargeError,
    TooManyAttachmentsError,
    UnsupportedAttachmentError,
    build_attachment_context,
    classify_attachment,
    downscale_image_for_context,
    extract_document_text,
    fetch_attachment_bytes,
    validate_attachment_set,
)


def _png_bytes(width: int, height: int) -> bytes:
    image = Image.new("RGB", (width, height), color=(200, 50, 50))
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return buf.getvalue()


def _client(handler) -> httpx.Client:
    return httpx.Client(transport=httpx.MockTransport(handler))


def test_classify_attachment_image_ok() -> None:
    assert classify_attachment("image/jpeg", 1024) == "image"


def test_classify_attachment_document_ok() -> None:
    assert classify_attachment("application/pdf", 1024) == "document"


def test_classify_attachment_unsupported_mime_raises() -> None:
    with pytest.raises(UnsupportedAttachmentError):
        classify_attachment("application/x-msdownload", 1024)


def test_classify_attachment_too_large_raises() -> None:
    with pytest.raises(AttachmentTooLargeError):
        classify_attachment("image/jpeg", 11 * 1024 * 1024)


def test_classify_attachment_video_size_cap_is_independent_of_image() -> None:
    # 15MB is over the image cap (10MB) but under the video cap (20MB)
    assert classify_attachment("video/mp4", 15 * 1024 * 1024) == "video"


def test_extract_document_text_plain() -> None:
    text = extract_document_text("Привет, мир".encode(), "text/plain")
    assert text == "Привет, мир"


def test_extract_document_text_markdown() -> None:
    text = extract_document_text(b"# Title\n\nBody text", "text/markdown")
    assert "Title" in text and "Body text" in text


def test_extract_document_text_docx() -> None:
    buf = io.BytesIO()
    document = Document()
    document.add_paragraph("Первый абзац.")
    document.add_paragraph("Второй абзац.")
    document.save(buf)

    text = extract_document_text(
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "Первый абзац." in text
    assert "Второй абзац." in text


def test_extract_document_text_pdf() -> None:
    fake_page = MagicMock()
    fake_page.extract_text.return_value = "Текст со страницы PDF"
    with patch("app.content_pipeline.attachments.PdfReader") as MockReader:
        MockReader.return_value.pages = [fake_page]
        text = extract_document_text(b"%PDF-fake-bytes", "application/pdf")
    assert text == "Текст со страницы PDF"


def test_extract_document_text_caps_length() -> None:
    text = extract_document_text(("x" * 20000).encode(), "text/plain")
    assert len(text) == 8000


def test_extract_document_text_collapses_whitespace() -> None:
    text = extract_document_text(
        "Первая строка.\n\n\n\n\nВторая  строка   с   пробелами.\n\n\n".encode(), "text/plain"
    )
    assert text == "Первая строка.\n\nВторая строка с пробелами."


def test_downscale_image_shrinks_oversized_image() -> None:
    original = _png_bytes(2000, 1000)
    resized, mime_type = downscale_image_for_context(original)
    assert mime_type == "image/jpeg"
    assert len(resized) < len(original)
    with Image.open(io.BytesIO(resized)) as image:
        # Both dimensions <= 384px is Gemini's one documented flat-rate
        # tile bucket (258 tokens regardless of aspect ratio) -- a
        # larger cap like 768 doesn't actually save anything, since the
        # tile-count formula scales with the image itself above that
        # point (verified by hand, see attachments.py).
        assert image.size == (384, 192)
        assert max(image.size) <= 384


def test_downscale_image_rejects_corrupt_data() -> None:
    with pytest.raises(UnsupportedAttachmentError):
        downscale_image_for_context(b"not-actually-an-image")


def test_downscale_image_leaves_small_image_dimensions_unchanged() -> None:
    original = _png_bytes(200, 100)
    resized, mime_type = downscale_image_for_context(original)
    assert mime_type == "image/jpeg"
    with Image.open(io.BytesIO(resized)) as image:
        assert image.size == (200, 100)


def test_extract_document_text_unsupported_mime_raises() -> None:
    with pytest.raises(UnsupportedAttachmentError):
        extract_document_text(b"data", "image/jpeg")


def test_fetch_attachment_bytes() -> None:
    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=b"raw-bytes")

    assert fetch_attachment_bytes("https://r2.example/x.jpg", client=_client(handler)) == b"raw-bytes"


def test_build_attachment_context_document_extracts_text() -> None:
    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=b"Some plain text content")

    context = build_attachment_context(
        "https://r2.example/notes.txt", "document", client=_client(handler)
    )
    assert context == {"kind": "text", "text": "Some plain text content"}


def test_build_attachment_context_image_returns_raw_bytes() -> None:
    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, content=b"fake-image-bytes")

    context = build_attachment_context(
        "https://r2.example/photo.png", "image", client=_client(handler)
    )
    assert context == {"kind": "media", "mime_type": "image/png", "data": b"fake-image-bytes"}


def test_validate_attachment_set_allows_up_to_five_images() -> None:
    validate_attachment_set(["image"] * 5)


def test_validate_attachment_set_allows_up_to_five_documents() -> None:
    validate_attachment_set(["document"] * 5)


def test_validate_attachment_set_allows_mixed_combination() -> None:
    # 3 photos + 1 document + 1 audio -- from the user's own example.
    validate_attachment_set(["image", "image", "image", "document", "audio"])


def test_validate_attachment_set_rejects_more_than_five_total() -> None:
    with pytest.raises(TooManyAttachmentsError):
        validate_attachment_set(["image"] * 6)


def test_validate_attachment_set_rejects_two_videos() -> None:
    with pytest.raises(TooManyAttachmentsError):
        validate_attachment_set(["video", "video"])


def test_validate_attachment_set_rejects_two_audios() -> None:
    with pytest.raises(TooManyAttachmentsError):
        validate_attachment_set(["audio", "audio"])


def test_validate_attachment_set_allows_one_video_and_one_audio_together() -> None:
    validate_attachment_set(["video", "audio", "image", "document"])


def test_validate_attachment_set_allows_empty() -> None:
    validate_attachment_set([])
